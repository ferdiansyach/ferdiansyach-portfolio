import os
import sys
import win32com.client
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Define CV Data
CONTACT = {
    "location": "Depok, Indonesia",
    "email": "iyanferdiansyach30@gmail.com",
    "phone": "+62 888 6007 599",
    "linkedin": "linkedin.com/in/ferdiansyach",
    "github": "github.com/ferdiansyach",
    "portfolio": "ferdiansyach-portfolio.vercel.app"
}

ROLES = {
    "generalist": {
        "title": {
            "id": "Information Systems Graduate | Versatile IT Professional | Web, Data & IT Support",
            "en": "Information Systems Graduate | Versatile IT Professional | Web, Data & IT Support"
        },
        "summary": {
            "id": "Lulusan Sistem Informasi (IPK 3.77) dengan pengalaman langsung di pengembangan web, analisis data, dan IT support. Membangun 3+ aplikasi web produksi (React, Next.js, Node.js) sekaligus merancang model prediktif ML (LSTM, XGBoost) dengan akurasi 92% di Telkom Indonesia. Terampil mengelola infrastruktur IT laboratorium dengan tingkat ketersediaan 98%, serta berpengalaman dalam metodologi Agile/Scrum. Mampu beradaptasi cepat di berbagai peran teknologi dan siap memberikan kontribusi lintas fungsi di lingkungan kerja yang dinamis.",
            "en": "Information Systems graduate (GPA 3.77) with hands-on experience spanning web development, data analysis, and IT support. Built 3+ production web applications (React, Next.js, Node.js) while engineering predictive ML models (LSTM, XGBoost) achieving 92% accuracy at Telkom Indonesia. Skilled in managing IT lab infrastructure with 98% device availability, and well-practiced in Agile/Scrum methodologies. A fast-adapting professional ready to contribute across diverse technology roles in dynamic work environments."
        }
    },
    "general": {
        "title": {
            "id": "IT Specialist | Quality Assurance & Manual Testing | Software Developer",
            "en": "IT Specialist | Quality Assurance & Manual Testing | Software Developer"
        },
        "summary": {
            "id": "Lulusan Sistem Informasi dengan kompetensi luas di bidang Quality Assurance (QA/Testing), rekayasa perangkat lunak, dan IT support. Terbukti memiliki ketelitian tinggi dalam mendeteksi bug dan mengoptimalkan performa sistem, termasuk merancang pipeline pengujian (Jest/RTL) yang memangkas bug rate sebesar 60% dan mempertahankan ketersediaan unit lab komputer sebesar 98%. Menguasai metodologi SDLC (Agile/Scrum), pengujian manual (API/Web), serta analisis data. Siap berkontribusi secara fleksibel di berbagai peran teknologi.",
            "en": "Information Systems graduate with broad competencies in Quality Assurance (QA/Testing), software engineering, and IT support. Proven track record of high attention to detail in bug detection and system optimization, including engineering a testing pipeline (Jest/RTL) that slashed production bug rate by 60% and maintaining 98% device availability in computer labs. Well-versed in SDLC (Agile/Scrum) methodologies, manual testing (API/Web), and data analysis. Ready to contribute flexibly across diverse IT roles."
        }
    },
    "fullstack": {
        "title": {
            "id": "Full-Stack Developer (React/Next.js, Node.js) | ML for Energy Forecasting",
            "en": "Full-Stack Developer (React/Next.js, Node.js) | ML for Energy Forecasting"
        },
        "summary": {
            "id": "Full-Stack Developer yang mengkhususkan diri dalam aplikasi berbasis data. Merancang dan meluncurkan 3+ aplikasi web produksi (React, Next.js, Node.js) dengan skor Lighthouse 90+. Merancang pipeline data end-to-end dan model ML prediktif (LSTM, XGBoost) mencapai akurasi 92% pada 50.000+ data poin di Telkom Indonesia. Menggabungkan keahlian web development dan data engineering untuk membangun produk perangkat lunak yang intelligent dan scalable.",
            "en": "Full-Stack Developer specializing in data-informed applications. Architected and shipped 3+ production web applications (React, Next.js, Node.js) with Lighthouse 90+ performance. Engineered end-to-end data pipelines and predictive ML models (LSTM, XGBoost) achieving 92% accuracy on 50,000+ data points at Telkom Indonesia. Leverages both web development and data engineering expertise to build intelligent, scalable software products."
        }
    },
    "data": {
        "title": {
            "id": "Data Analyst & ML Developer (Python, SQL) | Web-GIS & Predictive Modeling",
            "en": "Data Analyst & ML Developer (Python, SQL) | Web-GIS & Predictive Modeling"
        },
        "summary": {
            "id": "Data Analyst & Machine Learning Developer dengan keahlian kuat dalam Python, SQL, dan analisis geospasial. Merancang pipeline data end-to-end dan model ML prediktif (LSTM, XGBoost) mencapai akurasi 92% pada 50.000+ data poin di Telkom Indonesia. Berpengalaman membangun sistem monitoring kualitas air pesisir menggunakan Google Earth Engine dengan uji statistik Mann-Kendall. Ahli dalam mengubah dataset mentah dan time-series menjadi insight bisnis yang actionable melalui dashboard interaktif (Streamlit).",
            "en": "Data Analyst & Machine Learning Developer with strong expertise in Python, SQL, and geospatial analysis. Engineered end-to-end data pipelines and predictive ML models (LSTM, XGBoost) achieving 92% accuracy on 50,000+ data points at Telkom Indonesia. Experienced in building coastal water quality monitoring systems using Google Earth Engine with Mann-Kendall statistical tests. Adept at transforming raw, time-series datasets into actionable business insights through interactive dashboards (Streamlit)."
        }
    }
}

# The experience bullets remain 100% identical to the source React data, fully intact!
EXPERIENCES = [
    {
        "role": {"id": "Business Intelligence & ML Intern", "en": "Business Intelligence & ML Intern"},
        "company": "Telkom Indonesia | Infranexia",
        "type": "Internship",
        "location": "Jakarta, ID",
        "period": "Apr 2025 – Sep 2025",
        "bullets": [
            {
                "id": "Mengembangkan model AI prediktif menggunakan LSTM dan Gradient Boosting yang mencapai akurasi 92% dalam peramalan konsumsi energi gedung pintar, mendukung langsung divisi analisis tekno-ekonomi.",
                "en": "Engineered a predictive AI model using LSTM and Gradient Boosting that achieved 92% accuracy in smart building energy consumption forecasting, directly supporting the techno-economic analysis division."
            },
            {
                "id": "Merancang pipeline data end-to-end dari ingest mentah hingga deployment model (LSTM, XGBoost), memproses 50.000+ data poin dan menghasilkan insight yang actionable untuk optimasi efisiensi energi.",
                "en": "Architected an end-to-end data pipeline from raw ingestion through preprocessing to model deployment (LSTM, XGBoost), processing 50,000+ data points and generating actionable insights for energy efficiency optimization."
            },
            {
                "id": "Mengembangkan platform manajemen magang berbasis web bekerja sama dengan 3 divisi lintas fungsi, melayani 100+ peserta magang dan mengurangi overhead koordinasi onboarding.",
                "en": "Developed a web-based internship management platform in collaboration with 3 cross-functional divisions, serving 100+ interns and reducing onboarding coordination overhead."
            }
        ]
    },
    {
        "role": {"id": "Asisten Laboratorium Komputer", "en": "Computer Laboratory Assistant"},
        "company": "Lab Data Monetize | Universitas Nasional",
        "type": "Student Organization",
        "location": "Jakarta, ID",
        "period": "Sep 2024 – Sep 2025",
        "bullets": [
            {
                "id": "Memelihara dan men-troubleshoot 30+ unit komputer laboratorium, mempertahankan ketersediaan perangkat sebesar 98% selama jam operasional dan meminimalkan downtime sesi untuk 200+ mahasiswa per semester.",
                "en": "Maintained and troubleshot 30+ laboratory computer units, sustaining 98% device availability during operational hours and minimizing session downtime for 200+ students per semester."
            },
            {
                "id": "Mengoptimalkan alur instalasi perangkat lunak dan konfigurasi sistem, mendukung kegiatan praktikum 200+ mahasiswa per semester melalui prosedur pemeliharaan terstandarisasi.",
                "en": "Optimized software installation and system configuration workflows, supporting lab sessions for 200+ students per semester through standardized maintenance procedures."
            }
        ]
    },
    {
        "role": {"id": "Asisten Riset & Pengembangan", "en": "Research & Development Assistant"},
        "company": "HIMASI | Universitas Nasional",
        "type": "Student Organization",
        "location": "Jakarta, ID",
        "period": "Jun 2024 – Mar 2025",
        "bullets": [
            {
                "id": "Mengorkestrasi seminar teknologi nasional sebagai Ketua Pelaksana, mengoordinasikan 15 anggota panitia dari perencanaan hingga pelaksanaan, dan berhasil menghadirkan 150+ peserta tepat waktu.",
                "en": "Orchestrated a national technology seminar as Chief Executive, coordinating 15 committee members from planning to execution and successfully delivering the event to 150+ attendees on schedule."
            },
            {
                "id": "Mengembangkan dan mengelola website resmi himpunan menggunakan WordPress, meningkatkan traffic website sebesar 40% dan mendorong keterlibatan anggota melalui konten terstruktur yang dioptimalkan SEO.",
                "en": "Developed and managed the association's official WordPress website, increasing website traffic by 40% and driving member engagement through structured, SEO-optimized content."
            }
        ]
    },
    {
        "role": {"id": "Fullstack Developer", "en": "Fullstack Developer"},
        "company": "UNAS FEST | Universitas Nasional Festival",
        "type": "Student Organization",
        "location": "Jakarta, ID",
        "period": "Apr 2024 – Nov 2024",
        "bullets": [
            {
                "id": "Mengembangkan 10+ komponen website responsif menggunakan TypeScript dan Tailwind CSS, mencapai skor Lighthouse 90+ untuk performa dan aksesibilitas pada portal festival resmi.",
                "en": "Developed 10+ responsive website components using TypeScript and Tailwind CSS, achieving a Lighthouse score of 90+ for performance and accessibility on the official festival portal."
            },
            {
                "id": "Membangun API Routes Next.js untuk integrasi registrasi peserta dan database, serta merancang pipeline pengujian dengan Jest dan React Testing Library yang memangkas bug rate produksi sebesar 60%.",
                "en": "Built Next.js API Routes for attendee registration and database integration, and engineered a testing pipeline with Jest and React Testing Library that slashed production bug rate by 60%."
            }
        ]
    }
]

EDUCATION = {
    "institution": "Universitas Nasional",
    "degree": {
        "id": "Sarjana Komputer (S.Kom) | Sistem Informasi",
        "en": "Bachelor of Computer Science (S.Kom) | Information Systems"
    },
    "period": "2022 – 2026",
    "gpa": "3.77",
    "capstone": {
        "id": "Cloud-Based Coastal Water Quality Monitoring: Analisis Tren Spatiotemporal dan Unsupervised Spectral Clustering via Web-GIS",
        "en": "Cloud-Based Coastal Water Quality Monitoring: Spatiotemporal Trend Analysis and Unsupervised Spectral Clustering via Web-GIS"
    },
    "courses": {
        "id": "Pemrograman Web, Basis Data, Rekayasa Perangkat Lunak, Kecerdasan Buatan, Analisis & Perancangan Sistem, Jaringan Komputer",
        "en": "Web Programming, Database Systems, Software Engineering, Artificial Intelligence, System Analysis & Design, Computer Networks"
    }
}

TECHNICAL_SKILLS = {
    "generalist": {
        "id": [
            {"title": "Pengembangan Web", "skills": "HTML/CSS, JavaScript, TypeScript, Tailwind CSS, React, Next.js, Node.js, Express.js, REST API, GraphQL, Prisma, WordPress"},
            {"title": "Analisis Data & ML", "skills": "Python, NumPy, Pandas, Matplotlib, Scikit-learn, TensorFlow, NLTK, Jupyter Notebook, Streamlit, Tableau, Google Earth Engine, GIS, spatial analysis, clustering, unsupervised learning"},
            {"title": "Basis Data & DevOps", "skills": "MySQL / SQL, PostgreSQL, MongoDB, Docker, GCP, Git & GitHub, GitHub Actions, CI/CD"},
            {"title": "Metodologi & Soft Skills", "skills": "Agile Scrum, SDLC, Jira, Manual & API Testing, Postman, Dokumentasi Teknis, Kerja Tim Lintas Fungsi"},
            {"title": "IT Support & Administrasi", "skills": "Pemeliharaan Lab, Konfigurasi Sistem, Administrasi Windows & Linux, Troubleshooting Hardware/OS, Dasar Jaringan, Microsoft Office"}
        ],
        "en": [
            {"title": "Web Development", "skills": "HTML/CSS, JavaScript, TypeScript, Tailwind CSS, React, Next.js, Node.js, Express.js, REST API, GraphQL, Prisma, WordPress"},
            {"title": "Data Analysis & ML", "skills": "Python, NumPy, Pandas, Matplotlib, Scikit-learn, TensorFlow, NLTK, Jupyter Notebook, Streamlit, Tableau, Google Earth Engine, GIS, spatial analysis, clustering, unsupervised learning"},
            {"title": "Database & DevOps", "skills": "MySQL / SQL, PostgreSQL, MongoDB, Docker, GCP, Git & GitHub, GitHub Actions, CI/CD"},
            {"title": "Methodology & Soft Skills", "skills": "Agile Scrum, SDLC, Jira, Manual & API Testing, Postman, Technical Documentation, Cross-functional Teamwork"},
            {"title": "IT Support & Administration", "skills": "Lab Maintenance, System Configuration, Windows & Linux Administration, Hardware/OS Troubleshooting, Networking Basics, Microsoft Office"}
        ]
    },
    "general": {
        "id": [
            {"title": "Pengujian & QA (Testing)", "skills": "Manual Testing, API Testing (Postman), Unit & Component Testing (Jest & RTL), Web Developer Tools"},
            {"title": "Bahasa & Framework", "skills": "HTML/CSS, TypeScript, Python, React, Next.js, Node.js, Express.js, GIS, spatial analysis, clustering, unsupervised learning"},
            {"title": "Basis Data & Tools", "skills": "MySQL / SQL, MongoDB, Git & GitHub, GitHub Actions, Agile Scrum"},
            {"title": "IT Support & Troubleshooting", "skills": "Pemeliharaan Lab Komputer, Konfigurasi Sistem, Troubleshooting Hardware & OS"}
        ],
        "en": [
            {"title": "Software Testing & QA", "skills": "Manual Testing, API Testing (Postman), Unit & Component Testing (Jest & RTL), Web Developer Tools"},
            {"title": "Languages & Frameworks", "skills": "HTML/CSS, TypeScript, Python, React, Next.js, Node.js, Express.js, GIS, spatial analysis, clustering, unsupervised learning"},
            {"title": "Database & Tools", "skills": "MySQL / SQL, MongoDB, Git & GitHub, GitHub Actions, Agile Scrum"},
            {"title": "IT Support & Troubleshooting", "skills": "Lab Maintenance, System Configuration, Hardware & OS Troubleshooting"}
        ]
    },
    "fullstack": {
        "id": [
            {"title": "Frontend", "skills": "React, Next.js, TypeScript, Tailwind CSS, HTML / CSS"},
            {"title": "Backend & DB", "skills": "Prisma, Node.js, Express.js, MongoDB, MySQL / SQL, Postman"},
            {"title": "Data & Analisis", "skills": "Python, Pandas, NumPy, Matplotlib, Tableau, Streamlit, GIS, spatial analysis"},
            {"title": "Machine Learning", "skills": "Scikit-learn, XGBoost, LSTM / Keras, TensorFlow, clustering, unsupervised learning"},
            {"title": "DevOps & Tools", "skills": "Jest & RTL, GitHub Actions, Git & GitHub, Google Cloud, Docker, Jupyter Notebook"}
        ],
        "en": [
            {"title": "Frontend", "skills": "React, Next.js, TypeScript, Tailwind CSS, HTML / CSS"},
            {"title": "Backend & DB", "skills": "Prisma, Node.js, Express.js, MongoDB, MySQL / SQL, Postman"},
            {"title": "Data & Analysis", "skills": "Python, Pandas, NumPy, Matplotlib, Tableau, Streamlit, GIS, spatial analysis"},
            {"title": "Machine Learning", "skills": "Scikit-learn, XGBoost, LSTM / Keras, TensorFlow, clustering, unsupervised learning"},
            {"title": "DevOps & Tools", "skills": "Jest & RTL, GitHub Actions, Git & GitHub, Google Cloud, Docker, Jupyter Notebook"}
        ]
    },
    "data": {
        "id": [
            {"title": "Data & Analisis", "skills": "Python, Pandas, NumPy, Matplotlib, Tableau, Streamlit, GIS, spatial analysis"},
            {"title": "Machine Learning", "skills": "Scikit-learn, XGBoost, LSTM / Keras, TensorFlow, clustering, unsupervised learning"},
            {"title": "Backend & DB", "skills": "Prisma, Node.js, Express.js, MongoDB, MySQL / SQL, Postman"},
            {"title": "Frontend", "skills": "React, Next.js, TypeScript, Tailwind CSS, HTML / CSS"},
            {"title": "DevOps & Tools", "skills": "Jest & RTL, GitHub Actions, Git & GitHub, Google Cloud, Docker, Jupyter Notebook"}
        ],
        "en": [
            {"title": "Data & Analysis", "skills": "Python, Pandas, NumPy, Matplotlib, Tableau, Streamlit, GIS, spatial analysis"},
            {"title": "Machine Learning", "skills": "Scikit-learn, XGBoost, LSTM / Keras, TensorFlow, clustering, unsupervised learning"},
            {"title": "Backend & DB", "skills": "Prisma, Node.js, Express.js, MongoDB, MySQL / SQL, Postman"},
            {"title": "Frontend", "skills": "React, Next.js, TypeScript, Tailwind CSS, HTML / CSS"},
            {"title": "DevOps & Tools", "skills": "Jest & RTL, GitHub Actions, Git & GitHub, Google Cloud, Docker, Jupyter Notebook"}
        ]
    }
}

PROJECTS_DATA = {
    "indosaji": {
        "title": "Indosaji E-commerce",
        "description": {
            "id": "Platform e-commerce full-stack end-to-end berbasis MERN dengan integrasi payment gateway.",
            "en": "An end-to-end full-stack e-commerce platform based on MERN with payment gateway integration."
        },
        "technologies": ["React", "Node.js", "Express.js", "MongoDB", "REST API", "JWT Auth", "Stripe"],
        "githubUrl": "https://github.com/ferdiansyach/indosaji",
        "category": "webdev"
    },
    "smart-meter": {
        "title": "Smart Meter Analysis",
        "description": {
            "id": "Model AI prediktif (LSTM & XGBoost) untuk deteksi anomali energi dengan dashboard interaktif real-time.",
            "en": "Predictive AI models (LSTM & XGBoost) for energy anomaly detection with a real-time interactive dashboard."
        },
        "technologies": ["Python", "Machine Learning", "Data Analysis", "Streamlit", "XGBoost", "LSTM", "Pandas"],
        "githubUrl": "https://github.com/ferdiansyach/smart-meter-analysis",
        "category": "datascience"
    },
    "coastal-water-quality": {
        "title": "Coastal Water Quality Monitoring",
        "description": {
            "id": "Monitoring kualitas perairan pesisir Jakarta & Banten via Google Earth Engine (2019–2025) dengan Sentinel-2, Landsat-8, K-Means, dan Mann-Kendall.",
            "en": "Coastal water quality monitoring for Jakarta & Banten via Google Earth Engine (2019–2025) using Sentinel-2, Landsat-8, K-Means, and Mann-Kendall."
        },
        "technologies": ["Python", "Google Earth Engine", "Sentinel-2", "Landsat-8", "Streamlit", "K-Means", "Mann-Kendall", "Remote Sensing"],
        "githubUrl": "https://github.com/ferdiansyach/coastal-water-quality-gee",
        "category": "datascience"
    },
    "unasfest": {
        "title": "Website UNAS FEST",
        "description": {
            "id": "Portal website resmi berskala besar untuk mendukung operasional festival tahunan universitas.",
            "en": "A large-scale official web portal to support the operations of the university's annual festival."
        },
        "technologies": ["TypeScript", "React", "Tailwind CSS", "Next.js"],
        "githubUrl": "https://github.com/ferdiansyach/unasfest-end",
        "category": "webdev"
    },
    "himasi": {
        "title": "HIMASI UNAS Website",
        "description": {
            "id": "Sistem Manajemen Konten (CMS) profesional yang dioptimalkan untuk organisasi himpunan mahasiswa.",
            "en": "A professional Content Management System (CMS) optimized for the student association."
        },
        "technologies": ["WordPress", "Content Management", "SEO", "Plugin Management"],
        "githubUrl": "https://github.com/ferdiansyach/himasi-unas",
        "category": "wordpress"
    }
}

CERTIFICATIONS = [
    {
        "name": {"id": "Pengembang Web Pratama (Junior Web Developer)", "en": "Junior Web Developer"},
        "issuer": "BNSP (Badan Nasional Sertifikasi Profesi)",
        "date": "Feb 2026"
    },
    {
        "name": {"id": "IT Specialist | Python", "en": "IT Specialist | Python"},
        "issuer": "Certiport (Pearson VUE)",
        "date": "Oct 2025"
    },
    {
        "name": {"id": "Scrum Mastery", "en": "Scrum Mastery"},
        "issuer": "Telkom Indonesia",
        "date": "Jul 2025"
    },
    {
        "name": {"id": "Internet of Things Fundamental", "en": "Internet of Things Fundamental"},
        "issuer": "Telkom Indonesia",
        "date": "Apr 2025"
    },
    {
        "name": {"id": "Statistics Using R for Data Science", "en": "Statistics Using R for Data Science"},
        "issuer": "DQLab",
        "date": "Feb 2023"
    }
]

TRANSLATIONS = {
    "summary_title": {"id": "PROFIL PROFESIONAL", "en": "PROFESSIONAL SUMMARY"},
    "experience_title": {"id": "PENGALAMAN PROFESIONAL", "en": "PROFESSIONAL EXPERIENCE"},
    "education_title": {"id": "PENDIDIKAN", "en": "EDUCATION"},
    "skills_title": {"id": "KOMPETENSI TEKNIS", "en": "TECHNICAL SKILLS"},
    "projects_title": {"id": "PROYEK UNGGULAN", "en": "SELECTED PROJECTS"},
    "certifications_title": {"id": "SERTIFIKASI & LISENSI", "en": "CERTIFICATIONS & LICENSES"},
    "gpa_label": {"id": "IPK:", "en": "GPA:"},
    "capstone_label": {"id": "Proyek Capstone:", "en": "Capstone Project:"},
    "course_label": {"id": "Mata Kuliah Relevan:", "en": "Relevant Coursework:"},
}

def set_paragraph_spacing(p, before=0, after=1, line_spacing=1.0):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing

def add_p_border_bottom(p, color_hex="1E293B", size_eighth_pts=12):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size_eighth_pts))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_section_header(doc, text, lang, line_spacing=0.90):
    # Using built-in Heading 1 style
    p = doc.add_paragraph(style='Heading 1')
    set_paragraph_spacing(p, before=5.0, after=0.8, line_spacing=0.90)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(10.0)
    run.font.color.rgb = RGBColor(30, 41, 59) # #1e293b
    add_p_border_bottom(p, color_hex="334155", size_eighth_pts=8)

def generate_resume(role, lang, output_path, experiences_list):
    doc = Document()
    
    # Page setup - A4 Portrait with 0.5 inch margins
    sections = doc.sections
    for s in sections:
        s.page_width = Inches(8.27)
        s.page_height = Inches(11.69)
        s.top_margin = Inches(0.375)
        s.bottom_margin = Inches(0.375)
        s.left_margin = Inches(0.5)
        s.right_margin = Inches(0.5)
        
    # Configure built-in Heading 1 style to prevent Word default formatting
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Arial'
    h1_style.font.size = Pt(10.0)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(30, 41, 59)
    
    # Standard Paragraph style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9.5) # Body font: minimal 9.5-10pt
    font.color.rgb = RGBColor(55, 65, 81) # #374151
    
    # Line spacing: uniform 0.90 across all roles (matching verified fixed file)
    line_spacing = 0.90
    
    # 1. HEADER
    p_name = doc.add_paragraph()
    set_paragraph_spacing(p_name, before=0, after=0, line_spacing=line_spacing)
    run_name = p_name.add_run("FERDIANSYACH")
    run_name.bold = True
    run_name.font.name = 'Arial'
    run_name.font.size = Pt(16.5)
    run_name.font.color.rgb = RGBColor(15, 23, 42) # #0f172a
    
    # Header Subtitle
    p_sub = doc.add_paragraph()
    set_paragraph_spacing(p_sub, before=0, after=0, line_spacing=line_spacing)
    run_sub = p_sub.add_run(ROLES[role]["title"][lang])
    run_sub.bold = True
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(9.5)
    run_sub.font.color.rgb = RGBColor(29, 78, 216) # #1d4ed8
    
    # Contact row
    p_con = doc.add_paragraph()
    set_paragraph_spacing(p_con, before=0, after=0.2, line_spacing=0.90)
    contact_parts = [
        CONTACT["location"],
        CONTACT["email"],
        CONTACT["phone"],
        CONTACT["linkedin"],
        CONTACT["github"],
        CONTACT["portfolio"]
    ]
    contact_text = " | ".join(contact_parts)
    run_con = p_con.add_run(contact_text)
    run_con.font.name = 'Arial'
    run_con.font.size = Pt(8.5)
    run_con.font.color.rgb = RGBColor(71, 85, 105) # #475569
    
    # Line below header
    p_line = doc.add_paragraph()
    set_paragraph_spacing(p_line, before=0, after=0.2, line_spacing=0.1)
    run_line = p_line.add_run()
    run_line.font.size = Pt(1)
    add_p_border_bottom(p_line, color_hex="1e293b", size_eighth_pts=16) # 2pt line
    
    # 2. PROFESSIONAL SUMMARY
    add_section_header(doc, TRANSLATIONS["summary_title"][lang], lang, line_spacing)
    p_sum = doc.add_paragraph()
    set_paragraph_spacing(p_sum, before=0, after=1.0, line_spacing=line_spacing)
    p_sum.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_sum = p_sum.add_run(ROLES[role]["summary"][lang])
    run_sum.font.name = 'Arial'
    run_sum.font.size = Pt(9.5)
    
    # 3. PROFESSIONAL EXPERIENCE
    add_section_header(doc, TRANSLATIONS["experience_title"][lang], lang, line_spacing)
    
    # Per-role experience ordering (Isu 3 tailoring audit)
    # Fullstack & Data: Telkom, UNAS FEST, HIMASI, Lab Assistant
    # General (Manual Testing) & Generalist: chronological (as-is)
    import copy
    ordered_experiences = copy.deepcopy(experiences_list)
    if role in ("fullstack", "data", "general"):
        if role == "general":
            # Reorder: Telkom, UNAS FEST, Lab Assistant, HIMASI
            ordered_experiences = [
                copy.deepcopy(experiences_list[0]),  # Telkom
                copy.deepcopy(experiences_list[3]),  # UNAS FEST
                copy.deepcopy(experiences_list[1]),  # Lab Assistant
                copy.deepcopy(experiences_list[2]),  # HIMASI
            ]
        else:
            # Reorder: Telkom, UNAS FEST, HIMASI, Lab Assistant
            ordered_experiences = [
                copy.deepcopy(experiences_list[0]),  # Telkom
                copy.deepcopy(experiences_list[3]),  # UNAS FEST (Fullstack Developer)
                copy.deepcopy(experiences_list[2]),  # HIMASI (R&D Assistant)
                copy.deepcopy(experiences_list[1]),  # Lab Assistant
            ]
    
    # Fullstack only: reorder Telkom bullets - move web platform bullet (index 2) to first
    if role == "fullstack":
        telkom = ordered_experiences[0]
        if len(telkom["bullets"]) == 3:
            telkom["bullets"] = [telkom["bullets"][2], telkom["bullets"][0], telkom["bullets"][1]]
            
    # General (Manual Testing) only: split UNAS FEST bullet 2 and reframe Lab Assistant bullets
    if role == "general":
        for exp in ordered_experiences:
            if exp["company"] == "UNAS FEST | Universitas Nasional Festival":
                exp["bullets"] = [
                    exp["bullets"][0],  # bullet 1 tetap
                    {
                        "id": "Membangun API Routes Next.js untuk integrasi registrasi peserta dan database pada portal festival resmi.",
                        "en": "Built Next.js API Routes for attendee registration and database integration on the official festival portal."
                    },
                    {
                        "id": "Merancang dan mengeksekusi pipeline pengujian unit/komponen menggunakan Jest dan React Testing Library, menulis test case pada alur registrasi kritikal yang mengidentifikasi dan mencegah bug berulang — memangkas bug rate produksi sebesar 60% sebelum festival diluncurkan.",
                        "en": "Designed and executed a component/unit testing pipeline using Jest and React Testing Library, writing test cases across critical registration flows that identified and prevented recurring bugs — slashing production bug rate by 60% prior to festival launch."
                    }
                ]
            if exp["company"] == "Lab Data Monetize | Universitas Nasional":
                exp["bullets"] = [
                    {
                        "id": "Mendiagnosis dan menyelesaikan masalah hardware/software berulang pada 30+ unit lab melalui troubleshooting sistematis dan verifikasi pasca-perbaikan, mempertahankan ketersediaan perangkat 98% dan meminimalkan downtime sesi untuk 200+ mahasiswa per semester.",
                        "en": "Diagnosed and resolved recurring hardware/software issues across 30+ lab units through systematic troubleshooting and post-fix verification, sustaining 98% device availability and minimizing session downtime for 200+ students per semester."
                    },
                    {
                        "id": "Mendokumentasikan dan menstandarisasi prosedur instalasi/konfigurasi software menjadi checklist pemeliharaan yang dapat diulang, mengurangi kesalahan konfigurasi berulang pada sesi lab untuk 200+ mahasiswa per semester.",
                        "en": "Documented and standardized software installation/configuration procedures into repeatable maintenance checklists, reducing recurring configuration errors across lab sessions for 200+ students per semester."
                    }
                ]
    
    for exp in ordered_experiences:
        p_exp = doc.add_paragraph()
        # Spacing before new entry: 2.5pt (fully compliant with the 2-3pt requirement)
        set_paragraph_spacing(p_exp, before=2.5, after=0, line_spacing=line_spacing)
        
        # Tab stops: 7.27 inches is the right margin for 0.5 margin on A4
        p_exp.paragraph_format.tab_stops.add_tab_stop(Inches(7.27), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        
        run_role = p_exp.add_run(exp["role"][lang])
        run_role.bold = True
        run_role.font.name = 'Arial'
        run_role.font.size = Pt(10.0)
        run_role.font.color.rgb = RGBColor(15, 23, 42) # #0f172a
        
        p_exp.add_run("\t")
        
        run_period = p_exp.add_run(exp["period"])
        run_period.font.name = 'Arial'
        run_period.font.size = Pt(9.0)
        run_period.font.color.rgb = RGBColor(100, 116, 139) # #64748b
        
        # Company | Type | Location
        p_comp = doc.add_paragraph()
        set_paragraph_spacing(p_comp, before=0, after=0.1, line_spacing=line_spacing)
        run_comp = p_comp.add_run(exp["company"])
        run_comp.bold = True
        run_comp.font.name = 'Arial'
        run_comp.font.size = Pt(9.5)
        run_comp.font.color.rgb = RGBColor(29, 78, 216) # #1d4ed8
        
        exp_type_translated = {
            "Internship": "Magang" if lang == "id" else "Internship",
            "Student Organization": "Organisasi Mahasiswa" if lang == "id" else "Student Organization",
            "Freelance": "Lepas" if lang == "id" else "Freelance",
            "Full-time": "Penuh Waktu" if lang == "id" else "Full-time",
            "Part-time": "Paruh Waktu" if lang == "id" else "Part-time",
        }.get(exp["type"], exp["type"])
        
        run_meta = p_comp.add_run(f" | {exp_type_translated} | {exp['location']}")
        run_meta.font.name = 'Arial'
        run_meta.font.size = Pt(9.5)
        run_meta.font.color.rgb = RGBColor(71, 85, 105) # #475569
        
        # Bullets
        for bullet in exp["bullets"]:
            p_bullet = doc.add_paragraph(style='List Bullet')
            set_paragraph_spacing(p_bullet, before=0, after=0.1, line_spacing=line_spacing)
            p_bullet.paragraph_format.left_indent = Inches(0.25)
            run_bullet = p_bullet.add_run(bullet[lang])
            run_bullet.font.name = 'Arial'
            run_bullet.font.size = Pt(9.5)
            
    # 4. EDUCATION
    add_section_header(doc, TRANSLATIONS["education_title"][lang], lang, line_spacing)
    p_edu = doc.add_paragraph()
    set_paragraph_spacing(p_edu, before=2.5, after=0, line_spacing=line_spacing)
    p_edu.paragraph_format.tab_stops.add_tab_stop(Inches(7.27), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    
    run_inst = p_edu.add_run(EDUCATION["institution"])
    run_inst.bold = True
    run_inst.font.name = 'Arial'
    run_inst.font.size = Pt(10.0)
    run_inst.font.color.rgb = RGBColor(15, 23, 42) # #0f172a
    
    p_edu.add_run("\t")
    
    run_edu_period = p_edu.add_run(EDUCATION["period"])
    run_edu_period.font.name = 'Arial'
    run_edu_period.font.size = Pt(9.0)
    run_edu_period.font.color.rgb = RGBColor(100, 116, 139) # #64748b
    
    p_deg = doc.add_paragraph()
    set_paragraph_spacing(p_deg, before=0, after=0.2, line_spacing=line_spacing)
    run_deg = p_deg.add_run(EDUCATION["degree"][lang])
    run_deg.bold = True
    run_deg.font.name = 'Arial'
    run_deg.font.size = Pt(9.5)
    run_deg.font.color.rgb = RGBColor(30, 41, 59) # #1e293b
    
    # GPA
    p_gpa = doc.add_paragraph()
    set_paragraph_spacing(p_gpa, before=0, after=0.1, line_spacing=line_spacing)
    run_gpa_lbl = p_gpa.add_run(TRANSLATIONS["gpa_label"][lang] + " ")
    run_gpa_lbl.bold = True
    run_gpa_lbl.font.name = 'Arial'
    run_gpa_val = p_gpa.add_run(EDUCATION["gpa"])
    run_gpa_val.font.name = 'Arial'
    
    # Capstone
    p_cap = doc.add_paragraph()
    set_paragraph_spacing(p_cap, before=0, after=0.1, line_spacing=line_spacing)
    run_cap_lbl = p_cap.add_run(TRANSLATIONS["capstone_label"][lang] + " ")
    run_cap_lbl.bold = True
    run_cap_lbl.font.name = 'Arial'
    run_cap_val = p_cap.add_run(EDUCATION["capstone"][lang])
    run_cap_val.font.name = 'Arial'
    
    # Coursework
    p_crs = doc.add_paragraph()
    set_paragraph_spacing(p_crs, before=0, after=0.3, line_spacing=line_spacing)
    run_crs_lbl = p_crs.add_run(TRANSLATIONS["course_label"][lang] + " ")
    run_crs_lbl.bold = True
    run_crs_lbl.font.name = 'Arial'
    run_crs_val = p_crs.add_run(EDUCATION["courses"][lang])
    run_crs_val.font.name = 'Arial'
    
    # 5. TECHNICAL SKILLS
    add_section_header(doc, TRANSLATIONS["skills_title"][lang], lang, line_spacing)
    skills_list = TECHNICAL_SKILLS[role][lang]
    for skill_cat in skills_list:
        p_skill = doc.add_paragraph()
        set_paragraph_spacing(p_skill, before=0, after=0.1, line_spacing=line_spacing)
        run_title = p_skill.add_run(skill_cat["title"] + ": ")
        run_title.bold = True
        run_title.font.name = 'Arial'
        run_title.font.size = Pt(9.5)
        run_title.font.color.rgb = RGBColor(15, 23, 42)
        
        run_val = p_skill.add_run(skill_cat["skills"])
        run_val.font.name = 'Arial'
        run_val.font.size = Pt(9.5)
        
    # 6. SELECTED PROJECTS
    add_section_header(doc, TRANSLATIONS["projects_title"][lang], lang, line_spacing)
    
    selected_slugs = []
    if role == "generalist":
        selected_slugs = ["indosaji", "smart-meter", "himasi"]
    elif role == "general":
        selected_slugs = ["unasfest", "indosaji", "smart-meter"]
    elif role == "data":
        selected_slugs = ["smart-meter", "coastal-water-quality", "indosaji"]
    else: # fullstack
        selected_slugs = ["indosaji", "unasfest", "smart-meter"]
        
    for slug in selected_slugs:
        proj = PROJECTS_DATA[slug]
        p_proj = doc.add_paragraph()
        set_paragraph_spacing(p_proj, before=2.5, after=0, line_spacing=line_spacing)
        p_proj.paragraph_format.tab_stops.add_tab_stop(Inches(7.27), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        
        run_proj_title = p_proj.add_run(proj["title"])
        run_proj_title.bold = True
        run_proj_title.font.name = 'Arial'
        run_proj_title.font.size = Pt(10.0)
        run_proj_title.font.color.rgb = RGBColor(15, 23, 42)
        
        if proj.get("githubUrl"):
            run_git = p_proj.add_run(" (GitHub)")
            run_git.bold = True
            run_git.font.name = 'Arial'
            run_git.font.size = Pt(8.5)
            run_git.font.color.rgb = RGBColor(29, 78, 216)
            
        p_desc = doc.add_paragraph()
        set_paragraph_spacing(p_desc, before=0, after=0.1, line_spacing=line_spacing)
        run_desc = p_desc.add_run(proj["description"][lang])
        run_desc.font.name = 'Arial'
        run_desc.font.size = Pt(9.5)
        
        p_tech = doc.add_paragraph()
        set_paragraph_spacing(p_tech, before=0, after=0.3, line_spacing=line_spacing)
        run_tech = p_tech.add_run(", ".join(proj["technologies"]))
        run_tech.font.name = 'Arial'
        run_tech.font.size = Pt(8.5)
        run_tech.font.color.rgb = RGBColor(100, 116, 139)
        
    # 7. CERTIFICATIONS & LICENSES (Single Column)
    add_section_header(doc, TRANSLATIONS["certifications_title"][lang], lang, line_spacing)
    for cert in CERTIFICATIONS:
        p_cert = doc.add_paragraph()
        set_paragraph_spacing(p_cert, before=0, after=0.1, line_spacing=line_spacing)
        
        run_cert_name = p_cert.add_run(cert["name"][lang])
        run_cert_name.bold = True
        run_cert_name.font.name = 'Arial'
        run_cert_name.font.size = Pt(9.5)
        run_cert_name.font.color.rgb = RGBColor(15, 23, 42)
        
        run_cert_meta = p_cert.add_run(f" | {cert['issuer']} ({cert['date']})")
        run_cert_meta.font.name = 'Arial'
        run_cert_meta.font.size = Pt(9.5)
        
    # Save the document
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

def get_page_count_and_export_pdf(word_app, docx_path, pdf_path=None):
    doc = word_app.Documents.Open(os.path.abspath(docx_path))
    pages = doc.ComputeStatistics(2) # 2 = wdStatisticPages
    if pdf_path:
        os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17) # 17 = PDF
    doc.Close(False)
    return pages

if __name__ == "__main__":
    roles = ["generalist", "general", "fullstack", "data"]
    langs = ["id", "en"]
    output_dir = "public/cv"
    
    role_file_map = {
        "generalist": "Generalist",
        "general": "Manual_Testing",
        "fullstack": "Fullstack",
        "data": "Data"
    }
    
    # Initialize Word
    print("Opening Word COM application...")
    word_app = win32com.client.Dispatch("Word.Application")
    word_app.Visible = False
    
    try:
        for role in roles:
            for lang in langs:
                filename_base = f"Ferdiansyach_CV_{role_file_map[role]}_{lang.upper()}"
                docx_filename = f"{filename_base}.docx"
                pdf_filename = f"{filename_base}.pdf"
                docx_path = os.path.join(output_dir, docx_filename)
                pdf_path = os.path.join(output_dir, pdf_filename)
                
                print(f"Generating resume for {filename_base}...")
                generate_resume(role, lang, docx_path, EXPERIENCES)
                
                # Check page count and save PDF directly
                pages = get_page_count_and_export_pdf(word_app, docx_path, pdf_path)
                print(f"--> Generated: {docx_path} ({pages} pages)")
                if pages > 1:
                    print(f"WARNING: {filename_base} is {pages} pages!")
                    
    finally:
        print("Closing Word COM application...")
        word_app.Quit()
    
    # ==============================
    # POST-GENERATION REGRESSION CHECKS
    # ==============================
    print("\nRunning post-generation regression checks...")
    import zipfile
    import xml.etree.ElementTree as ET
    
    regression_errors = []
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    for role in roles:
        for lang in langs:
            filename_base = f"Ferdiansyach_CV_{role_file_map[role]}_{lang.upper()}"
            docx_path = os.path.join(output_dir, f"{filename_base}.docx")
            
            with zipfile.ZipFile(docx_path) as z:
                with z.open('word/document.xml') as f:
                    raw = f.read().decode('utf-8')
            
            # Check 1: "5+" should not appear in any summary
            if '5+' in raw and ('5+ production' in raw.lower() or '5+ aplikasi' in raw.lower()):
                regression_errors.append(f"{filename_base}: Still contains '5+' in summary (should be '3+')")
            
            # Check 2: Redis should not appear in fullstack/data skills
            if role in ("fullstack", "data") and 'Redis' in raw:
                regression_errors.append(f"{filename_base}: Still contains 'Redis' in skills")
            
            # Check 3: Figma should not appear in fullstack/data skills
            if role in ("fullstack", "data") and 'Figma' in raw:
                regression_errors.append(f"{filename_base}: Still contains 'Figma' in skills")
    
    if regression_errors:
        print("\n[FAIL] REGRESSION CHECK FAILED:")
        for err in regression_errors:
            print(f"  [X] {err}")
        sys.exit(1)
    else:
        print("[OK] All regression checks passed!")
    
    print("Done!")
